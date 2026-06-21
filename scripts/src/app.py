"""
--- Stakeholder Gap Analysis DESIGN NOTES---
PURPOSE: A tool for auditing stakeholder engagement in project management.

PIPELINE: 1) ingest documents, 2) extract stakeholder-related facts, normalize
those facts into a common schema, retrieve evidence from multiple sources,
run gap-detection rules or prompts, then produce a grounded output with citations.

--- RETRIEVAL LAYER ---

1) INGESTION: The first step is document ingestion. You pull in the core
artifacts you care about:  stakeholder register, stakeholder engagement
plan, communications management plan, issue log, RAID log if relevant,
meeting notes, steering committee minutes, and status updates. Early on,
keep this narrow. A beginner-friendly version might only use four sources:
stakeholder register, engagement plan, issue log, and meeting notes.

2) CHUNKING: The second step is document parsing and chunking. You
convert each source into machine-readable text, then split it into chunks.
For ordinary RAG, chunking is about semantic retrieval. For this app,
chunking also needs document awareness. Each chunk should carry metadata
such as document type, project name, date, stakeholder names mentioned,
meeting type, and section heading. That metadata
will matter a lot later when you compare facts across documents.

3) EXTRACTION: The third step is entity extraction and normalization.
This is one of the most important parts. You need to identify stakeholder
names, teams, roles, concerns, influence indicators, communication cadences,
decisions, issues, action items, and sentiment clues. Then you normalize them.
For example, “Finance Director,” “Head of Finance,” and “Maria Chen” may
all refer to the same stakeholder. If you skip normalization, your app will
think they are separate people and your gap detection will be noisy.

A useful mental model is that you are building a lightweight stakeholder
knowledge graph or at least a structured table of facts. Instead of only
storing raw chunks in a vector database, you also store extracted facts in
structured form. For example:

stakeholder: Maria Chen
role: Finance Director
source: steering committee notes
concern: budget approval timing
date: 2026-05-12
sentiment: concerned
action owner: PM
And separately:

stakeholder: Maria Chen
listed_in_register: false

That structured layer makes comparison much easier than relying only on free-text
retrieval.

4) VECTOR DATABASE: The fourth step is embedding and indexing. You embed
the chunks and store them in a vector database such as FAISS or Chroma. This
supports semantic retrieval later. At the same time, you store normalized entities and facts
in a relational store like SQLite or Postgres. In practice, this means your
\app becomes a hybrid system: vector search for evidence retrieval, structured
queries for consistency checking. [Will do both in JSON]

5) MAPPING: The fifth step is artifact-to-schema mapping. Each document type contributes
different facts. The stakeholder register gives you stakeholder identity, influence,
interest, current engagement level, and desired engagement level. The engagement
plan gives you strategies and actions. The communications plan gives you cadence,
channel, and audience. Meeting notes give you observed behavior, concerns, attendance,
objections, and decisions. The issue log gives you active pain points and unresolved
themes. You map all of these into common comparison fields so the system can ask
questions like: “This concern appears in notes and issues; where is the corresponding
engagement action?”

6) RETRIEVAL: The sixth step is retrieval orchestration. When the app runs a gap
check, it should not just do one similarity search. It should retrieve by stakeholder,
by concern, by time window, and by artifact type. For example, if checking whether
a stakeholder is missing from the register, the app retrieves meeting-note chunks
mentioning that stakeholder, issue-log entries tied to that stakeholder, and the
stakeholder register entry or absence of one. This is more targeted than generic RAG.

7) GAP DETECTION: The seventh step is gap detection logic. This is the core of the
app. I would not rely only on an LLM here. Use a combination of deterministic rules
and LLM reasoning.

Deterministic rules work well for things like:

stakeholder mentioned more than three times in notes but absent from the register
stakeholder marked high influence in register but missing communication cadence
recurring concern appears in notes across two or more weeks but no related issue
or engagement action exists
stakeholder register says neutral, but recent notes suggest escalating resistance
communication plan says monthly update, but no evidence of outreach in six weeks

The LLM then helps with fuzzier checks such as:
whether repeated meeting comments represent the same underlying concern
whether two differently worded role labels likely refer to the same stakeholder
whether a concern has been addressed in an engagement plan even if the wording differs
This hybrid approach is much more reliable than pure prompting.

8) EVIDENCE GROUNDING: The eighth step is evidence grounding and scoring. Every
detected gap should include supporting evidence and a confidence score. For example:

Candidate gap: Finance Director is materially engaged but absent from stakeholder register.
Evidence: Mentioned in 4 meeting notes and 2 issue entries over 5 weeks.
Confidence: High.

Or:

Candidate gap: Communications cadence for Operations sponsor may be insufficient.
Evidence: Register shows high influence; communications plan lists quarterly updates;
two recent notes mention requests for weekly visibility.
Confidence: Medium.

This matters because stakeholder data is messy, and you want the tool to propose
candidate inconsistencies, not pretend it knows ground truth.

9) REPORT GENERATION: The ninth step is output generation. The app should produce
a concise, reviewable report, not just a chat answer. A strong output format is
one row per detected gap with fields such as gap type, stakeholder, severity,
evidence summary, source links, and suggested PM action. The LLM can generate
a narrative summary on top of that, but the structured findings should come first.

10) HITL: The tenth step is human review loop. This is important in stakeholder
management because names, politics, and informal influence are often ambiguous.
The PM should be able to mark a finding as confirmed, false positive, or already
addressed. Those review decisions can become feedback data to improve rules and prompts.

MVP pipeline might look like this:

1) INGESTION: Upload four artifact types.
2) CHUNKING: Parse text and split into chunks with metadata.
3a) EXTRACTION: Extract stakeholder names, roles, concerns, cadence, and engagement signals.
3b) NORMALIZATION: Normalize stakeholder identities.
4) VECTOR DATABASE: Store text chunks in a vector DB and extracted facts in SQL.
6a) Run scheduled checks for a small set of predefined gaps.
6b) Retrieve supporting passages for each gap.
9) Use an LLM to summarize the inconsistency and explain why it matters.
9) Show the PM a candidate-gap dashboard with citations.
10) Let the PM confirm or dismiss each gap.

One design caution: do not let the model invent stakeholder facts. Make the app
cite sources and separate observed evidence from recommendations. For example,
“Observed: Legal raised data retention concerns in two meetings. Recommendation:
schedule targeted follow-up.” That distinction will make the app more trustworthy.

The highest-risk part technically is entity resolution. If your app cannot tell
whether “Steering Sponsor,” “Program Sponsor,” and “Alicia” are the same person,
the rest gets weak fast. So for a beginner, I would simplify aggressively: start
with a manually maintained stakeholder alias file. That lets you avoid building
sophisticated identity matching too early.

The best MVP gap checks are probably these three:

First, missing stakeholder detection. If someone appears often in notes or issues
but is not in the stakeholder register, flag it.

Second, missing engagement coverage. If a high-influence stakeholder exists in the
register but has no clear engagement strategy or communication cadence, flag it.

Third, recurrent concern mismatch. If a concern appears repeatedly in notes or
the issue log but is not reflected in the engagement plan, flag it.

Those three alone would make a credible first product.

A simple retrieval pattern for one gap might look like this:

query SQL for stakeholders seen in notes in the last 30 days
compare that list to stakeholders in the register
for each missing stakeholder, retrieve top note chunks mentioning them
ask the LLM to summarize why this person appears to be a relevant stakeholder
generate a candidate finding with evidence citations
That is much easier to build than a fully open-ended chat app.

One more design principle: separate facts, inferences, and recommendations.

Facts are things like “mentioned in 4 meetings.”
Inferences are things like “likely influential.”
Recommendations are things like “add to register and define engagement strategy.”

If you keep those layers separate, the app will be more trustworthy and easier to debug.

For a beginner Python build, I would recommend this stack:

The biggest conceptual takeaway is this: a Stakeholder Gap Detector is not just
“ask a chatbot about documents.”
It is a retrieval + normalization + rule comparison + explanation pipeline.
RAG supports the evidence retrieval, but the app’s real value comes from comparing
structured signals across multiple project artifacts.

Here is the shortest version of the pipeline:

Ingest -> Chunk -> Extract facts -> Normalize stakeholders ->
Store text and facts -> Run gap rules -> Retrieve evidence ->
Use LLM to explain -> Present candidate gaps for PM review

A good next build step would be to define the schema first.
Start with four entities: Stakeholder, Concern, EngagementAction,
and CommunicationCadence. Then map each artifact into those entities
before touching the UI.
"""


import json
import os
from datetime import date
from pathlib import Path
import contextlib
import numpy as np
import pandas as pd
import streamlit as st
from openai import OpenAI
from pydantic import BaseModel, Field, ConfigDict
from typing import Annotated
import docx


# Define paths and global variables
# This gives you a datetime object
today_obj = date.today()
today = today_obj.strftime("%B %d, %Y")

src_dir = Path(__file__).resolve().parent
scripts_dir = src_dir.parent
project_root = scripts_dir.parent

log_folder = project_root / "logs"
output_folder = project_root / "outputs"
data_folder = project_root / "data"
vector_store_folder = project_root / "vector_store"

stakeholder_gap_report_path = output_folder / "STAKEHOLDER_GAP_REPORT.txt"
database_file_destination = vector_store_folder / "global_vector_store.json"

# Ensure the paths exists
folder_paths = [log_folder, output_folder, data_folder, vector_store_folder]
for folder in folder_paths:
    if not folder.exists():
        folder.mkdir(parents=True)


# Initialize LLM client
api_key = os.environ.get("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# Pydantic models to define the input and output formats.
NarrativeText = Annotated[
    str,
    Field(
        min_length=20,
        description="Write in 2 to 3 clear, professional sentences."
    )
]

# This is the model for the stakeholder gap report to be used in the stakeholder gap dashboard.

class StakeholderGapReport(BaseModel):
    gap_category: str
    stakeholder_name: str
    severity: str
    confidence: str
    observed_gap: str
    practical_impact: str
    recommended_action: str
    evidence: list[EvidenceItem]

    class EvidenceItem(BaseModel):
        source: str
        snippet: str

    gap_category: Annotated[
        str,
        Field(description="The stakeholder gap category in ALL CAPS, for example MISSING STAKEHOLDER or EXECUTION GAP.")
    ]
    observed_gap: Annotated[
        str,
        Field(description="A concise description of the stakeholder engagement gap.")
    ]
    practical_impact: Annotated[
        str,
        Field(description="The likely project or stakeholder impact if the gap is not addressed.")
    ]
    recommended_action: Annotated[
        str,
        Field(description="A concrete action the project manager should take to address the gap.")
    ]

# This is the model for the executive stakeholder gap report to be used in the stakeholder gap dashboard.
class ExecutiveStakeholderGapReport(BaseModel):
    model_config = ConfigDict(extra="forbid")

    executive_summary: Annotated[
        str,
        Field(description="A professional summary of the main stakeholder engagement findings in 2 to 3 sentences.")
    ]
    categories: Annotated[
        list[StakeholderGapReport],
        Field(min_length=1, description="Detailed stakeholder gap findings by category.")
    ]

# This class is used to capture stdout and redirect it to a Streamlit placeholder.
class StreamlitStdoutRedirector:
    def __init__(self, placeholder):
        self.placeholder = placeholder
        self.output_str = ""

    def write(self, text):
        self.output_str += text
        self.placeholder.code(self.output_str, language="text")

    def flush(self):
        pass

graph = {
    "nodes": [],
    "edges": []
}

def save_graph(graph: dict, filepath: Path) -> None:
    with filepath.open("w", encoding="utf-8") as f:
        json.dump(graph, f, indent=2)

def load_graph(filepath: Path) -> dict:
    if filepath.exists():
        with filepath.open("r", encoding="utf-8") as f:
            return json.load(f)
    return {"nodes": [], "edges": []}

# This function adds a node to the graph.
def add_node(graph, node_id, node_type, name):
    if not any(n["id"] == node_id for n in graph["nodes"]):
        graph["nodes"].append({
            "id": node_id,
            "type": node_type,
            "name": name
        })

# This function adds an edge to the graph.
def add_edge(graph, source, target, edge_type):
    graph["edges"].append({
        "source": source,
        "target": target,
        "type": edge_type
    })

# This function creates a graph object.
def create_graph():
    return {
        "nodes": [],
        "edges": []
    }

# This function builds a stakeholder graph
def build_stakeholder_graph(register_rows, meeting_records, engagement_rows):
    graph = create_graph()

    for row in register_rows:
        stakeholder_id = f"stakeholder_{row['name'].lower().replace(' ', '_')}"
        add_node(graph, stakeholder_id, "Stakeholder", row["name"])

    for meeting in meeting_records:
        meeting_id = f"meeting_{meeting['date']}_{meeting['title'].lower().replace(' ', '_')}"
        add_node(graph, meeting_id, "Meeting", meeting["title"])

        for attendee in meeting.get("attendees", []):
            stakeholder_id = f"stakeholder_{attendee.lower().replace(' ', '_')}"
            add_edge(graph, stakeholder_id, meeting_id, "ATTENDED")

        for concern in meeting.get("concerns", []):
            concern_id = f"concern_{concern.lower().replace(' ', '_')}"
            add_node(graph, concern_id, "Concern", concern)
            for speaker in meeting.get("attendees", []):
                stakeholder_id = f"stakeholder_{speaker.lower().replace(' ', '_')}"
                add_edge(graph, stakeholder_id, concern_id, "MENTIONED")

    return graph

# This function gets the edges for a given node.
def get_edges_for_node(graph, node_id):
    return [e for e in graph["edges"] if e["source"] == node_id or e["target"] == node_id]

# This function gets an embedding from the OpenAI API.
def get_embedding(text: str, model: str = "text-embedding-3-small") -> list[float]:
    cleaned_text = str(text).replace("", " ").strip() or ""
    response = client.embeddings.create(
        model=model,
        input=cleaned_text
    )
    return response.data[0].embedding

# This function calculates the cosine similarity between two vectors.
def cosine_similarity(v1: list[float], v2: list[float]) -> float:
    a = np.array(v1, dtype=float)
    b = np.array(v2, dtype=float)

    denominator = np.linalg.norm(a) * np.linalg.norm(b)
    if denominator == 0:
        return 0.0

    return float(np.dot(a, b) / denominator)

# This function chunks the .docx stakeholder plan using the docx library
def chunk_stakeholder_plan(filepath: Path) -> list[dict]:
    chunks = []
    filename = filepath.name
    doc = docx.Document(filepath)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return chunks

# This function chunks the .docx meeting notes using the docx library
def chunk_meeting_notes(filepath: Path) -> list[dict]:
    chunks = []
    filename = filepath.name
    doc = docx.Document(filepath)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return chunks

# This function chunks the .xlsx stakeholder register using the pandas library
def chunk_stakeholder_register(filepath: Path) -> list[dict]:
    chunks = []
    filename = filepath.name
    df = pd.read_excel(filepath)
    for index, row in df.iterrows():
        chunks.append({
            "text": row["Stakeholder Name"],
            "source": filename
        })
        chunks.append({
        "text": "This is the end of the register.",
            "source": filename
        })
    return chunks


# This function processes a folder and returns a list of chunks.
def process_folder(folder_path: Path, chunk_size: int = 500, overlap: int = 80) -> list[dict]:
    all_chunks = []

    if not folder_path.exists():
        print(f" -> Directory target folder '{folder_path}' missing. Indexer aborted.")
        return []

    for file_path in folder_path.iterdir():
        if file_path.is_dir():
            continue

        file_chunks = []
        file_name_lower = file_path.name.lower()
        file_suffix = file_path.suffix.lower()

        if file_name_lower == "Meeting_Notes.docx":
            print(f" -> Chunking Meeting Notes: {file_path.name}")
            file_chunks = chunk_meeting_notes(file_path)
        elif file_name_lower == "Stakeholder_Plan.docx":
            print(f" -> Chunking Stakeholder Engagement Strategy: {file_path.name}")
            file_chunks = chunk_stakeholder_plan(file_path, chunk_size, overlap)
        elif file_name_lower == "Stakeholder_Register.xlsx":
            print(f" -> Chunking Stakeholder Register: {file_path.name}")
            file_chunks = chunk_stakeholder_register(file_path)
        else:
            continue

        all_chunks.extend(file_chunks)

    return all_chunks

# This class is responsible for loading and storing the knowledge graph.
class StakeholderFactStore:
    def __init__(self, filepath: Path):
        self.filepath = filepath
        self.graph = {"nodes": [], "edges": []}
        self.load()

    def load(self) -> None:
        if self.filepath.exists():
            with self.filepath.open("r", encoding="utf-8") as f:
                self.graph = json.load(f)
        else:
            self.graph = {"nodes": [], "edges": []}

    def get_nodes_by_type(self, node_type: str) -> list[dict]:
        return [n for n in self.graph["nodes"] if n["type"] == node_type]

    def get_registered_stakeholders(self) -> set[str]:
        return {
            n["name"].strip().lower()
            for n in self.graph["nodes"]
            if n["type"] == "Stakeholder"
        }

# This class is responsible for detecting gaps in the knowledge graph.
class GapDetector:
    def __init__(self, fact_store):
        self.fact_store = fact_store

    def find_missing_stakeholders(self):
        mentioned = self.fact_store.get_stakeholders_mentioned_in_notes()
        registered = self.fact_store.get_registered_stakeholders()
        missing = [s for s in mentioned if s not in registered]
        return missing

    def find_missing_cadence(self):
        high_influence = self.fact_store.get_high_influence_stakeholders()
        no_cadence = [
            s for s in high_influence
            if not self.fact_store.has_fact(s, "HAS_COMMUNICATION_CADENCE")
        ]
        return no_cadence

    def find_unaddressed_concerns(self):
        concerns = self.fact_store.get_recurring_concerns()
        return [
            c for c in concerns
            if not self.fact_store.concern_has_engagement_action(c)
            ]

# This class is responsible for storing and retrieving chunks of text.
class SimpleVectorStore:
    def __init__(self):
        self.entries = []

    def load(self, filepath: Path) -> None:
        if filepath.exists():
            with filepath.open("r", encoding="utf-8") as f:
                self.entries = json.load(f)
            print(f" -> Vector Store loaded {len(self.entries)} entries.")
        else:
            raise FileNotFoundError(f" Vector store file not found.")

    def add_many(self, chunks: list[dict]) -> None:
        print(f" -> Sending {len(chunks)} chunks to OpenAI for vector synthesis...")
        for chunk in chunks:
            embedding = get_embedding(chunk["text"])
            self.entries.append({**chunk, "embedding": embedding})

    def search(self, query: str, top_k: int = 3) -> list[dict]:
        query_embedding = get_embedding(query)
        scored = []
        for entry in self.entries:
            sim = cosine_similarity(query_embedding, entry["embedding"])
            scored.append((sim, entry))

        scored.sort(reverse=True, key=lambda x: x[0])
        return [{**entry, "similarity": round(sim, 4)} for sim, entry in scored[:top_k]]

    def save(self, filepath: Path) -> None:
        print(f" -> Saving vector store entries.")
        with filepath.open("w", encoding="utf-8") as f:
            json.dump(self.entries, f, indent=4)
            print(f" -> Vector Store file saved as {filepath.name}")


# This class is responsible for matching stakeholder names to known aliases.
class StakeholderMatcher:
    def __init__(self, alias_map: dict[str, str] | None = None):
        self.alias_map = alias_map or {}

    def normalize_name(self, raw_name: str) -> str:
        normalized = raw_name.strip().lower()
        return self.alias_map.get(normalized, normalized)


def run_gap_analysis(fact_store: StakeholderFactStore, vector_store: SimpleVectorStore) -> list[dict]:
    detector = GapDetector(fact_store)

    findings = []

    for stakeholder in detector.find_missing_stakeholders():
        evidence = vector_store.search(f"{stakeholder} stakeholder meeting concern", top_k=3)
        findings.append({
            "gap_category": "MISSING STAKEHOLDER",
            "observed_gap": f"{stakeholder} appears in project evidence but is missing from the register.",
            "practical_impact": "An influential or active stakeholder may be unmanaged.",
            "recommended_action": "Review the stakeholder register and define an engagement approach.",
            "evidence": evidence
        })

    return findings



# --- REFACTORED REPORT GENERATION VIA PYDANTIC TARGETS ---

def generate_gap_report(
    structured_report: ExecutiveStakeholderGapReport,
    file_path: Path
) -> str:
    lines = []
    lines.append("STAKEHOLDER GAP ANALYSIS REPORT")
    lines.append("")
    lines.append("Executive Summary:")
    lines.append(structured_report.executive_summary)
    lines.append("")

    for item in structured_report.categories:
        lines.append(f"Category: {item.gap_category}")
        lines.append(f"Observed Gap: {item.observed_gap}")
        lines.append(f"Practical Impact: {item.practical_impact}")
        lines.append(f"Recommended Action: {item.recommended_action}")
        lines.append("")

    final_report_text = "\n".join(lines)

    with file_path.open("w", encoding="utf-8") as f:
        f.write(final_report_text)

    return final_report_text


# --- REFACTORED WORKFLOW ROUTER WITH PYDANTIC PARSING ---
def synthesize_report_with_llm(audit_results: list[dict]) -> ExecutiveStakeholderGapReport:
    """
    Takes raw unregistered risks and leverages OpenAI's structured output beta mechanics
    to return a strictly checked Pydantic model representation.
    """
    print(" -> Sending raw risk data to OpenAI for structural executive synthesis...")

    raw_context = json.dumps(audit_results, indent=2)

    prompt = f"""
    You are an expert Senior Project Leader. Analyze the following raw data of stakeholder gaps. 

    Raw Gap Data Data:
    {raw_context}
    """

    # We leverage beta.chat.completions.parse to enforce Pydantic parsing at the API wire layer
    response = client.beta.chat.completions.parse(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": "You are a precise, professional Senior Project Leader. Extract risk anomalies into the required schema structure cleanly without markdown text wraps.",
            },
            {"role": "user", "content": prompt},
        ],
        response_format=ExecutiveStakeholderGapReport,
        temperature=0.3,
    )

    # This returns the initialized ExecutiveRiskReport object directly
    return response.choices[0].message.parsed


# --- CORE PIPELINE EXECUTION WRAPPER ---
def run_automated_pipeline(log_placeholder):
    try:
        print("PIPELINE STARTED.")

        # --- 1. CREATING VECTOR STORE ---
        print("STEP 1: Creating Vector Store.")
        store = SimpleVectorStore()

        # 1. Ensure the directory path exists
        target_dir = database_file_destination.parent
        if not target_dir.exists():
            print(f" -> Creating missing directory: {target_dir}")
            target_dir.mkdir(parents=True, exist_ok=True)

        # 2. Check for the file
        if database_file_destination.exists():
            print(f" -> Found existing vector store: {database_file_destination.name}")
            store.load(database_file_destination)
        else:
            print(" -> No existing vector store found. Starting new ingestion...")
            compiled_data_chunks = process_folder(data_folder, chunk_size=500, overlap=80)

            if not compiled_data_chunks:
                print("No data found to process. Exiting.")
                return

            store.add_many(compiled_data_chunks)
            store.save(database_file_destination)

        # --- 2. Identifying unregistered risks ---
        print(f"STEP 2: Starting automated risk audit.")
        matcher = StakeholderMatcher(register_path) if register_path.exists() else None

        if not matcher:
            print(f"Warning: Register not found at {register_path}. Skipping registration check.")

        gap_queries = [
            "staffing turnover, resource departures, personnel shortages",
            "security vulnerabilities, mTLS failures, unauthorized data access",
            "data pipeline errors, system latency, memory leaks, parsing crashes"
        ]

        discovered_gap_data = run_gap_analysis(store, matcher, audit_queries)

        # --- 3. Synthesize with LLM ---
        print("STEP 3: Synthesizing AI Report Narrative via Structured Validation.")
        # Returns an actual Pydantic ExecutiveRiskReport object instance now
        structured_report_obj = synthesize_report_with_llm(discovered_gap_data)

        # --- 4. Generate file on disk ---
        print("STEP 4: Generating Structured Risk Audit Report Artifacts.")
        final_report_text = generate_gap_report(structured_report_obj, discovered_gap_data, report_path)

        print("PIPELINE COMPLETED.")
        return final_report_text

    except Exception as e:
        print(f"Pipeline crashed with an unhandled traceback exception: {e}")

    return None


# --- STREAMLIT UI CONFIGURATION ---
st.set_page_config(
    page_title="AI Stakeholder Gap Analysis",
    layout="wide"
)

st.title("Stakeholder Gap Analysis Dashboard")
st.markdown("---")

col1, col2 = st.columns(2)

with col1:
    st.subheader("System Configuration")
    target_directory = data_folder

    if target_directory.exists():
        st.text(f"Files found in '{target_directory.name}'")
        # iterdir() yields Path objects; we grab .name for just the filename
        files = [f.name for f in target_directory.iterdir()]
        st.write(files)
    else:
        st.error(f"Directory not found")

    start_pipeline = st.button("Generate Stakeholder Gap Analysis Report", use_container_width=True, type="primary")

    st.subheader("Pipeline Summary")
    console_logs = st.empty()
    console_logs.info("Click 'Generate Stakeholder Gap Analysis Report' button to begin.")

with col2:
    st.subheader("Report Workspace")
    report_placeholder = st.empty()
    report_placeholder.info("The Stakeholder Gap Analysis Report will populate here upon synthesis.")

if start_pipeline:
    redirector = StreamlitStdoutRedirector(console_logs)

    with st.spinner("Processing stakeholder parameters..."):
        with contextlib.redirect_stdout(redirector):
            final_narrative = run_automated_pipeline(console_logs)

    if final_narrative:
        with report_placeholder.container():
            st.html(
                f"""
                <div style="
                    background-color: #1e293b; 
                    color: #f8fafc; 
                    padding: 20px; 
                    border-radius: 8px; 
                    height: 550px; 
                    overflow-y: scroll; 
                    white-space: pre-wrap; 
                    font-family: inherit;
                    border: 1px solid #334155;
                    line-height: 1.5;
                ">
                    <p style="font-size: 16px !important; margin: 0; padding: 0;">{final_narrative}</p>
                </div>
                """
            )

            st.download_button(
                label="Download Stakeholder Gap Report (.txt)",
                data=final_narrative,
                file_name="audit_report.txt",
                mime="text/plain",
                use_container_width=True
            )